
import io
import docx
import pdfplumber
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd

from preprocessing import preprocess


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []

    # Обычные параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Таблицы — перебираем все ячейки
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Внутри ячейки тоже могут быть параграфы
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)

    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    parts = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:

            #  стандартное извлечение
            page_text = page.extract_text()
            if page_text and page_text.strip():
                parts.append(page_text)
                continue

            # собираем текст из отдельных слов
            # (помогает когда вёрстка многоколоночная или сложная)
            words = page.extract_words()
            if words:
                # Сортируем слова по вертикали (сверху вниз), потом по горизонтали
                words_sorted = sorted(words, key=lambda w: (round(w["top"] / 5), w["x0"]))
                page_text = " ".join(w["text"] for w in words_sorted)
                if page_text.strip():
                    parts.append(page_text)

    return "\n".join(parts)


def extract_text(uploaded_file) -> str:
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        return ""


def compute_similarity(texts: list[str], filenames: list[str], method: str) -> pd.DataFrame:
    processed = []
    for text in texts:
        tokens = preprocess(text, method=method)
        processed.append(" ".join(tokens))

    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(processed)
    sim_matrix = cosine_similarity(tfidf_matrix)

    sim_percent = (sim_matrix * 100).round(1)
    df = pd.DataFrame(sim_percent, index=filenames, columns=filenames)
    return df